from docx import Document
import mysql.connector
import string

# Replace these variables with your database information
host = ''
user = ''
password = ''
database = ''

# Establish a connection
conn = mysql.connector.connect(
    host=host,
    user=user,
    password=password,
    database=database
)

# Check if the connection was successful
if conn.is_connected():
    print("Successfully connected to the database")

# extract names from word doc
def extract_names(docx_file):
    doc = Document(docx_file)

    paragraphs = doc.paragraphs[:1]
    Names = []

    for paragraph in paragraphs:
        words = paragraph.text.split()
        for word in words:
            # Remove punctuation and add to Names list if it's not 'Name:'
            cleaned_word = word.strip(string.punctuation)
            if cleaned_word and 'Name' not in cleaned_word:
                Names.append(cleaned_word)
    return Names

# extract group number
def extract_group_number(docx_file):
    doc = Document(docx_file)

    paragraphs = doc.paragraphs[:10]

    for paragraph in paragraphs:
        words = paragraph.text.split()
        
        for word in words:
            if word.isdigit():
                return int(word)

        
doc_file = "Exercise8.docx"

names = extract_names(doc_file)
group_number = extract_group_number(doc_file)
missing = []

cursor = conn.cursor()

query = f"SELECT * FROM my_table WHERE Group_Number = {group_number};"
cursor.execute(query)

row = cursor.fetchall()

conn.close()

for i in range(len(row)):
    first_name = row[i][1]
    last_name = row[i][2]
    if first_name not in names and last_name not in names:
        missing.append((first_name, last_name))

def print_missing():
    if len(missing) == 0:
        print("No one is missing")
    else:
        print("Missing: ", end="")
        for i in range(len(missing)):
            print(missing[i][0], missing[i][1])

print_missing()